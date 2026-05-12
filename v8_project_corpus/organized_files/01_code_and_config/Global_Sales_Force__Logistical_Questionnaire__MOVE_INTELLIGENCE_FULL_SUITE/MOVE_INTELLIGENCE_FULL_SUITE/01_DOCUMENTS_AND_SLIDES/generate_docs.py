import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- EXCEL GENERATION ---
def create_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "Logistics Intake"
    
    # Theme: Elegant Black
    THEME = {
        'primary': '2D2D2D',
        'light': 'E5E5E5',
        'accent': '2D2D2D'
    }
    
    ws.sheet_view.showGridLines = False
    ws.column_dimensions['A'].width = 3
    
    # Title
    ws['B2'] = "[COMPANY NAME] - LOGISTICS INTAKE FORM"
    ws['B2'].font = Font(name='Source Serif Pro', size=18, bold=True, color=THEME['primary'])
    
    # Headers
    headers = ["Client Name", "Move Date", "Pick-Up Address", "PU Building Type", "PU Truck Access", "PU Stairs", "PU Elevator", "Delivery Address", "Del Building Type", "Del Truck Access", "Del Stairs", "Del Elevator", "Load Preference", "Specialty Items"]
    
    for col_num, header in enumerate(headers, 2):
        col_letter = get_column_letter(col_num)
        cell = ws[f'{col_letter}5']
        cell.value = header
        cell.font = Font(name='Source Serif Pro', size=11, bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color=THEME['primary'], end_color=THEME['primary'], fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.column_dimensions[col_letter].width = 20
        
    # Sample Data Row
    sample_data = ["John Doe", "2024-08-15", "123 Main St", "Single Family", "Clear (46ft+)", "None", "No", "456 High St", "Apartment", "Restricted", "2 Flights", "Yes (Reserved)", "Live Load", "Piano"]
    for col_num, data in enumerate(sample_data, 2):
        col_letter = get_column_letter(col_num)
        cell = ws[f'{col_letter}6']
        cell.value = data
        cell.font = Font(name='Source Sans Pro', size=11)
        cell.alignment = Alignment(horizontal='left', vertical='center')
        
    # Borders
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    for row in ws['B5:O6']:
        for cell in row:
            cell.border = thin_border
            
    wb.save('/home/ubuntu/logistics_suite/docs/Logistics_Intake_Template.xlsx')

# --- WORD DOC GENERATION ---
def create_word():
    doc = Document()
    
    # Title
    title = doc.add_heading('[COMPANY NAME] - Logistics Intake Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Welcome to the future of moving. Please review the logistics requirements below.")
    
    # Insert Images
    doc.add_heading('Level 1: Standard Access', level=1)
    doc.add_picture('/home/ubuntu/logistics_suite/assets/hero_normal_move.png', width=Inches(6))
    doc.add_paragraph("A standard move requires 46ft of clear space, no stairs, and no elevator restrictions.")
    
    doc.add_heading('Level 5: Expert Access', level=1)
    doc.add_picture('/home/ubuntu/logistics_suite/assets/hero_expert_move.png', width=Inches(6))
    doc.add_paragraph("Expert moves involve long carries, multiple flights of stairs, elevator reservations, and parking permits.")
    
    doc.add_heading('Truck Dimensions', level=1)
    doc.add_picture('/home/ubuntu/logistics_suite/assets/truck_diagram.png', width=Inches(6))
    doc.add_paragraph("Our 26ft box trucks require 46ft of total space. Semi-trailers require 80ft.")
    
    doc.save('/home/ubuntu/logistics_suite/docs/Logistics_Visual_Guide.docx')

if __name__ == "__main__":
    create_excel()
    create_word()
